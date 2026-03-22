import os
import uuid
import json
import io

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, FileResponse, Response
from pydantic import BaseModel
import anthropic
import stripe
import pdfplumber
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.enums import TA_CENTER
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

app = FastAPI(title="ResumeAI")

# In-memory store — swap for Redis/Postgres in production
submissions: dict = {}

stripe.api_key = os.environ.get("STRIPE_SECRET_KEY", "")
claude = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))

BASE_URL = os.environ.get("BASE_URL", "http://localhost:8000")
PRICE_CENTS = int(os.environ.get("PRICE_CENTS", "799"))  # $7.99


# ── Models ─────────────────────────────────────────────────────────────────────

class CheckoutRequest(BaseModel):
    resume_text: str
    job_description: str


class AnalyzeRequest(BaseModel):
    session_id: str


class DownloadRequest(BaseModel):
    text: str


# ── API routes ─────────────────────────────────────────────────────────────────

@app.post("/api/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """Extract plain text from an uploaded PDF resume."""
    if not file.filename.lower().endswith(".pdf"):
        raise HTTPException(400, "Only PDF files are supported")

    content = await file.read()
    try:
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            text = "\n".join(
                page.extract_text() or "" for page in pdf.pages
            ).strip()
        if not text:
            raise HTTPException(400, "Could not extract text from PDF — try pasting your resume instead")
        return {"text": text}
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(400, f"Failed to parse PDF: {e}")


@app.post("/api/checkout")
async def create_checkout(req: CheckoutRequest):
    """Store submission data and create a Stripe Checkout session."""
    if not req.resume_text.strip():
        raise HTTPException(400, "Resume text is required")
    if not req.job_description.strip():
        raise HTTPException(400, "Job description is required")

    submission_id = str(uuid.uuid4())
    submissions[submission_id] = {
        "resume_text": req.resume_text,
        "job_description": req.job_description,
        "result": None,
    }

    try:
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[
                {
                    "price_data": {
                        "currency": "usd",
                        "product_data": {
                            "name": "AI Resume Optimizer",
                            "description": "ATS score · optimized resume · tailored cover letter",
                        },
                        "unit_amount": PRICE_CENTS,
                    },
                    "quantity": 1,
                }
            ],
            mode="payment",
            success_url=f"{BASE_URL}/?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{BASE_URL}/",
            metadata={"submission_id": submission_id},
        )
    except stripe.error.StripeError as e:
        raise HTTPException(500, f"Stripe error: {e.user_message}")

    return {"url": session.url}


@app.post("/api/analyze")
async def analyze(req: AnalyzeRequest):
    """Verify payment then run Claude analysis. Results are cached per submission."""
    try:
        session = stripe.checkout.Session.retrieve(req.session_id)
    except stripe.error.InvalidRequestError:
        raise HTTPException(400, "Invalid session ID")

    if session.payment_status != "paid":
        raise HTTPException(402, "Payment not completed")

    submission_id = session.metadata.get("submission_id")
    if not submission_id or submission_id not in submissions:
        raise HTTPException(404, "Submission not found — it may have expired. Please start over.")

    sub = submissions[submission_id]

    # Return cached result on repeated calls (e.g., page refresh)
    if sub["result"]:
        return sub["result"]

    result = _run_claude_analysis(sub["resume_text"], sub["job_description"])
    sub["result"] = result
    return result


# ── Download endpoints ────────────────────────────────────────────────────────

@app.post("/api/download/docx")
async def download_docx(req: DownloadRequest):
    """Generate and return an optimized resume as a .docx file."""
    content = _create_docx(req.text)
    return Response(
        content=content,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=optimized-resume.docx"},
    )


@app.post("/api/download/pdf")
async def download_pdf(req: DownloadRequest):
    """Generate and return an optimized resume as a .pdf file."""
    content = _create_pdf(req.text)
    return Response(
        content=content,
        media_type="application/pdf",
        headers={"Content-Disposition": "attachment; filename=optimized-resume.pdf"},
    )


def _create_docx(text: str) -> bytes:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    lines = text.strip().split("\n")
    is_first = True

    for line in lines:
        s = line.strip()
        if not s:
            doc.add_paragraph("")
            continue

        if is_first:
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.bold = True
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            is_first = False
        elif s.isupper() and len(s) < 50:
            p = doc.add_paragraph()
            run = p.add_run(s)
            run.bold = True
            run.font.size = Pt(11)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "888888")
            pBdr.append(bottom)
            pPr.append(pBdr)
        else:
            doc.add_paragraph(s)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _create_pdf(text: str) -> bytes:
    # Sanitize characters that reportlab's built-in fonts can't handle
    replacements = {
        "\u2022": "-", "\u2014": "-", "\u2013": "-",
        "\u2019": "'", "\u2018": "'", "\u201c": '"', "\u201d": '"',
        "\u00e9": "e", "\u00e8": "e", "\u00ea": "e",
        "&": "&amp;", "<": "&lt;", ">": "&gt;",
    }
    for char, repl in replacements.items():
        text = text.replace(char, repl)

    name_style = ParagraphStyle("Name", fontName="Helvetica-Bold", fontSize=18,
                                alignment=TA_CENTER, spaceAfter=4)
    header_style = ParagraphStyle("Header", fontName="Helvetica-Bold", fontSize=11,
                                  spaceBefore=14, spaceAfter=3)
    body_style = ParagraphStyle("Body", fontName="Helvetica", fontSize=10.5,
                                leading=15, spaceAfter=2)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter,
                            leftMargin=inch, rightMargin=inch,
                            topMargin=inch, bottomMargin=inch)
    story = []
    lines = text.strip().split("\n")
    is_first = True

    for line in lines:
        s = line.strip()
        if not s:
            story.append(Spacer(1, 6))
            continue

        if is_first:
            story.append(Paragraph(s, name_style))
            is_first = False
        elif s.isupper() and len(s) < 50:
            story.append(Paragraph(s, header_style))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=colors.grey, spaceAfter=4))
        else:
            story.append(Paragraph(s, body_style))

    doc.build(story)
    buf.seek(0)
    return buf.read()


# ── Claude analysis ────────────────────────────────────────────────────────────

def _run_claude_analysis(resume_text: str, job_description: str) -> dict:
    prompt = f"""You are an expert resume coach and ATS (Applicant Tracking System) specialist.

Analyze the resume below against the job description and return a JSON object with EXACTLY this structure (no markdown, no extra keys):

{{
  "ats_score": <integer 0-100>,
  "score_explanation": "<2-3 sentence explanation of the score>",
  "missing_keywords": ["<keyword>", ...],
  "improvements": ["<specific actionable improvement>", ...],
  "optimized_resume": "<full rewritten resume in plain text>",
  "cover_letter": "<full tailored cover letter in plain text>"
}}

Rules:
- missing_keywords: up to 10 important terms/skills from the job description absent from the resume
- improvements: 5-8 specific, actionable bullet points (e.g. "Add 'Python' to Skills section")
- optimized_resume: complete rewrite preserving all true facts, improving wording and keyword density
- cover_letter: professional, 3-paragraph letter tailored to this specific role

RESUME:
{resume_text}

JOB DESCRIPTION:
{job_description}"""

    message = claude.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}],
    )

    raw = message.content[0].text
    # Strip markdown code fences if present
    raw = raw.strip()
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1]
        raw = raw.rsplit("```", 1)[0]

    start = raw.find("{")
    end = raw.rfind("}") + 1
    return json.loads(raw[start:end])


# ── Static files (must be last) ────────────────────────────────────────────────

app.mount("/", StaticFiles(directory="static", html=True), name="static")
